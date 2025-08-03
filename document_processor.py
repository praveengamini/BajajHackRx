import requests
import PyPDF2
import io
import tempfile
import os
import pickle
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document as LangchainDocument
from fastapi import HTTPException
from config import Config
import uuid

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE, 
            chunk_overlap=Config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""]
        )
        self.embeddings = HuggingFaceEmbeddings(model_name=Config.EMBEDDING_MODEL_NAME)
        
        # Create directories for FAISS indexes and metadata
        os.makedirs(Config.FAISS_INDEX_PATH, exist_ok=True)
        os.makedirs(Config.FAISS_METADATA_PATH, exist_ok=True)
    
    def download_document(self, url: str) -> bytes:
        """Download document from URL"""
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            return response.content
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to download document: {str(e)}")
    
    def extract_text_from_pdf(self, pdf_content: bytes) -> str:
        """Extract text from PDF content"""
        try:
            pdf_file = io.BytesIO(pdf_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to extract PDF text: {str(e)}")
    
    def extract_text_from_docx(self, docx_content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            with tempfile.NamedTemporaryFile() as tmp_file:
                tmp_file.write(docx_content)
                tmp_file.flush()
                doc = Document(tmp_file.name)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to extract DOCX text: {str(e)}")
    
    def process_document(self, document_url: str) -> tuple:
        """Process document and create FAISS vector store"""
        content = self.download_document(document_url)
        
        if document_url.lower().endswith('.pdf'):
            text = self.extract_text_from_pdf(content)
        elif document_url.lower().endswith('.docx'):
            text = self.extract_text_from_docx(content)
        else:
            text = content.decode('utf-8', errors='ignore')
        
        # Split text into chunks
        text_chunks = self.text_splitter.split_text(text)
        
        # Create Langchain documents
        documents = [LangchainDocument(page_content=chunk) for chunk in text_chunks]
        
        # Create FAISS vector store
        vector_store = FAISS.from_documents(documents, self.embeddings)
        
        # Generate unique identifier for this document
        doc_id = f"doc_{str(uuid.uuid4())[:8]}"
        
        # Save FAISS index
        index_path = os.path.join(Config.FAISS_INDEX_PATH, doc_id)
        vector_store.save_local(index_path)
        
        # Save metadata
        metadata = {
            "doc_id": doc_id,
            "source": document_url,
            "num_chunks": len(text_chunks),
            "original_text": text
        }
        metadata_path = os.path.join(Config.FAISS_METADATA_PATH, f"{doc_id}.pkl")
        with open(metadata_path, 'wb') as f:
            pickle.dump(metadata, f)
        
        return vector_store, text, doc_id
    
    def load_vector_store(self, doc_id: str) -> FAISS:
        """Load existing FAISS vector store"""
        index_path = os.path.join(Config.FAISS_INDEX_PATH, doc_id)
        if not os.path.exists(index_path):
            raise HTTPException(status_code=404, detail=f"Vector store not found for doc_id: {doc_id}")
        
        vector_store = FAISS.load_local(index_path, self.embeddings, allow_dangerous_deserialization=True)
        return vector_store
    
    def get_metadata(self, doc_id: str) -> dict:
        """Get metadata for a document"""
        metadata_path = os.path.join(Config.FAISS_METADATA_PATH, f"{doc_id}.pkl")
        if not os.path.exists(metadata_path):
            return {}
        
        with open(metadata_path, 'rb') as f:
            return pickle.load(f)
    
    def list_documents(self) -> list:
        """List all processed documents"""
        if not os.path.exists(Config.FAISS_METADATA_PATH):
            return []
        
        doc_list = []
        for filename in os.listdir(Config.FAISS_METADATA_PATH):
            if filename.endswith('.pkl'):
                doc_id = filename[:-4]  # Remove .pkl extension
                metadata = self.get_metadata(doc_id)
                doc_list.append(metadata)
        
        return doc_list